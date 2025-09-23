import csv
import subprocess
import sys
import os

def get_user_preferences():
    """Get user preferences for what to include in the generated files."""
    print("\n=== ICT Exercise Generator Configuration ===")
    print("Please choose what to include in your exercise files:\n")
    
    # Show source information
    while True:
        show_source = input("Show source information (year, chapter, topic)? (y/n): ").lower().strip()
        if show_source in ['y', 'yes', 'n', 'no']:
            show_source_info = show_source in ['y', 'yes']
            break
        print("Please enter 'y' for yes or 'n' for no")
    
    # Show answers
    while True:
        show_answer = input("Show answers? (y/n): ").lower().strip()
        if show_answer in ['y', 'yes', 'n', 'no']:
            show_answers = show_answer in ['y', 'yes']
            break
        print("Please enter 'y' for yes or 'n' for no")
    
    # Show answer spaces (only ask if not showing answers)
    if not show_answers:
        while True:
            show_space = input("Show answer spaces for long questions? (y/n): ").lower().strip()
            if show_space in ['y', 'yes', 'n', 'no']:
                show_answer_spaces = show_space in ['y', 'yes']
                break
            print("Please enter 'y' for yes or 'n' for no")
    else:
        show_answer_spaces = False  # No need for spaces if showing answers
    
    preferences = {
        'show_source_info': show_source_info,
        'show_answers': show_answers,
        'show_answer_spaces': show_answer_spaces
    }
    
    print(f"\n📋 Your preferences:")
    print(f"  • Source info: {'Yes' if show_source_info else 'No'}")
    print(f"  • Show answers: {'Yes' if show_answers else 'No'}")
    print(f"  • Answer spaces: {'Yes' if show_answer_spaces else 'No'}")
    print()
    
    return preferences

def get_lq_answers():
    """Load LQ answers from CSV for display."""
    lq_answers = {}
    try:
        with open('LQ_Answers.csv', 'r', encoding='utf-8') as csv_file:
            reader = csv.DictReader(csv_file)
            for row in reader:
                question_id = row.get('QuestionIDwithQuestionyears', '')
                if question_id:
                    answers = {}
                    for i in range(1, 9):  # Up to 8 sub-answers
                        answer_key = f'SubAnswer{i}'
                        if answer_key in row and row[answer_key].strip():
                            answers[i] = row[answer_key].strip()
                    if answers:
                        lq_answers[question_id] = answers
    except FileNotFoundError:
        print("Note: LQ_Answers.csv not found. Long question answers will not be displayed.")
    except Exception as e:
        print(f"Error reading LQ answers: {e}")
    
    return lq_answers

def install_required_packages():
    """Install required packages if not already installed."""
    packages_to_install = []
    
    # Check for python-docx
    try:
        import docx
        print("✓ python-docx is already installed")
    except ImportError:
        packages_to_install.append('python-docx')
    
    # Install missing packages
    for package in packages_to_install:
        try:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package], 
                                capture_output=True, text=True)
            print(f"✓ Successfully installed {package}")
        except subprocess.CalledProcessError as e:
            print(f"✗ Failed to install {package}")
            print("Please try installing manually:")
            print(f"  pip3 install {package}")
            print(f"  or: python3 -m pip install {package}")
            return False
    
    return True

def generate_exercise_markdown(preferences):
    """Generates the Exercise.md file from MCQ.csv and LQ.csv with user preferences."""
    
    # Load LQ answers if needed
    lq_answers = get_lq_answers() if preferences['show_answers'] else {}
    
    try:
        with open('Exercise.md', 'w', encoding='utf-8') as md_file:
            # Write the YAML frontmatter
            md_file.write("---\n")
            md_file.write("marp: true\n")
            md_file.write("theme: UT\n")
            md_file.write("class: title-page\n")
            md_file.write("paginate: false\n")
            md_file.write("backgroundColor: white\n")
            md_file.write("---\n\n")

            # Write the title page
            title_suffix = ""
            if preferences['show_answers']:
                title_suffix = " (答案版)"
            elif not preferences['show_answer_spaces']:
                title_suffix = " (純題目版)"
                
            md_file.write("# **佛教黃鳳翎中學**\n\n")
            md_file.write(f"# **ICT 練習題集{title_suffix}**\n\n")
            md_file.write("---\n\n")

            # Write the Multiple-Choice (MCQ) section
            md_file.write('<div class="section-title"><strong>甲部 多項選擇題</strong></div>\n')
            md_file.write('<strong>請選擇最合適的答案。</strong>\n\n')

            # Process MCQ questions
            try:
                with open('MCQ.csv', 'r', encoding='utf-8') as csv_file:
                    reader = csv.DictReader(csv_file)
                    question_counter = 1
                    
                    for row in reader:
                        # Skip empty rows
                        if not row.get('QuestionText', '').strip():
                            continue
                        
                        # Extract source information
                        question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                        book_chapter = row.get('book_chapter', '')
                        topics = row.get('topics', '')
                        
                        # Parse year from QuestionIDwithQuestionyears (e.g., "2016-1-6" -> "2016")
                        year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                        
                        # Extract data from the CSV format
                        question_text = row['QuestionText']
                        option_a = row['OptionA']
                        option_b = row['OptionB']
                        option_c = row['OptionC']
                        option_d = row['OptionD']
                        marks = row['Marks']
                        correct_answer = row.get('answer', '').upper() if preferences['show_answers'] else ''
                        
                        # Add a divider after every 5 questions for visual break
                        if question_counter > 1 and question_counter % 5 == 1:
                            md_file.write('---\n\n')

                        # Write source information if enabled
                        if preferences['show_source_info']:
                            source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                            md_file.write(f'<div class="question-source">{source_info}</div><br>\n')
                        
                        md_file.write(f'<div class="question-item">{question_counter}. {question_text} <span class="points">({marks} 分)</span></div>\n')
                        md_file.write('<ul class="mcq-options">\n')
                        
                        # Write options with answer highlighting if enabled
                        options = [('A', option_a), ('B', option_b), ('C', option_c), ('D', option_d)]
                        for letter, option_text in options:
                            if preferences['show_answers'] and letter == correct_answer:
                                md_file.write(f'<li style="color: red; font-weight: bold;">{option_text} ✓</li>\n')
                            else:
                                md_file.write(f'<li>{option_text}</li>\n')
                        
                        md_file.write('</ul>\n\n')
                        question_counter += 1
                        
            except FileNotFoundError:
                print("Error: 'MCQ.csv' not found.")
                return False
            except KeyError as e:
                print(f"Error: Missing column in MCQ.csv: {e}")
                return False

            # Write the Long-Answer (LQ) section
            md_file.write('---\n\n')
            md_file.write('<div class="section-title"><strong>乙部 問答題</strong></div>\n')
            if preferences['show_answer_spaces']:
                md_file.write('<strong>請在適當的答案框內作答。</strong>\n\n')
            else:
                md_file.write('<strong>請回答以下問題。</strong>\n\n')

            # Process LQ questions
            try:
                with open('LQ.csv', 'r', encoding='utf-8') as csv_file:
                    reader = csv.DictReader(csv_file)
                    lq_counter = 1
                    
                    for row in reader:
                        # Skip empty rows
                        if not row.get('QuestionText', '').strip():
                            continue
                        
                        # Extract source information
                        question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                        book_chapter = row.get('book_chapter', '')
                        topics = row.get('topics', '')
                        
                        # Parse year and question ID
                        year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                        question_id = question_id_with_years.split('-')[-1] if question_id_with_years else str(lq_counter)
                        
                        question_text = row['QuestionText']
                        total_marks = row['Marks']
                        
                        md_file.write('---\n\n')
                        
                        # Write source information if enabled
                        if preferences['show_source_info']:
                            source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                            md_file.write(f'<div class="question-source">{source_info}</div><br>\n')
                        
                        md_file.write(f'<div class="question-item">{lq_counter}. {question_text} <span class="points">({total_marks} 分)</span></div>\n\n')

                        # Handle long questions with sub-questions (up to 8 sub-questions)
                        sub_questions = []
                        for i in range(1, 9):  # SubQuestion1 to SubQuestion8
                            sub_text = row.get(f'SubQuestion{i}', '') or ''
                            sub_marks = row.get(f'SubMarks{i}', '') or ''
                            sub_text = sub_text.strip() if sub_text else ''
                            sub_marks = sub_marks.strip() if sub_marks else ''
                            
                            if sub_text:
                                sub_questions.append({
                                    'text': sub_text,
                                    'marks': sub_marks,
                                    'index': i
                                })

                        # Display all sub-questions found
                        for idx, sub in enumerate(sub_questions):
                            md_file.write(f'<div class="sub-question-item">{sub["text"]} <span class="points">({sub["marks"]}分)</span></div>\n')
                            
                            # Show answers if enabled
                            if preferences['show_answers']:
                                # Look for answer in lq_answers
                                answer_text = ""
                                if question_id_with_years in lq_answers and sub['index'] in lq_answers[question_id_with_years]:
                                    answer_text = lq_answers[question_id_with_years][sub['index']]
                                elif question_id_with_years in lq_answers and (idx + 1) in lq_answers[question_id_with_years]:
                                    answer_text = lq_answers[question_id_with_years][idx + 1]
                                
                                if answer_text:
                                    md_file.write(f'<div style="color: red; font-weight: bold; margin: 10px 0; padding: 10px; border-left: 3px solid red;">答案: {answer_text}</div>\n')
                                else:
                                    md_file.write('<div style="color: orange; font-style: italic;">答案: [答案未提供]</div>\n')
                            
                            # Show answer spaces if enabled
                            elif preferences['show_answer_spaces']:
                                marks_value = int(sub["marks"]) if sub["marks"] and sub["marks"].isdigit() else 2
                                if marks_value >= 4:
                                    height = "5em"
                                elif marks_value >= 2:
                                    height = "4em"
                                else:
                                    height = "3em"
                                md_file.write(f'<div class="answer-box"><div style="height: {height};"></div></div>\n')

                            md_file.write('\n')
                        
                        lq_counter += 1

            except FileNotFoundError:
                print("Note: 'LQ.csv' not found. Skipping long questions section.")
            except KeyError as e:
                print(f"Error: Missing column in LQ.csv: {e}")

        print("Successfully generated Exercise.md!")
        return True
        
    except Exception as e:
        print(f"Error creating Exercise.md: {e}")
        return False

def generate_exercise_docx(preferences):
    """Generates the Exercise.docx file from MCQ.csv and LQ.csv with user preferences."""
    
    try:
        # Import docx modules here after installation
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Load LQ answers if needed
        lq_answers = get_lq_answers() if preferences['show_answers'] else {}
        
        # Create a new Document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        # Title page
        title = doc.add_heading('佛教黃鳳翎中學', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_suffix = ""
        if preferences['show_answers']:
            title_suffix = " (答案版)"
        elif not preferences['show_answer_spaces']:
            title_suffix = " (純題目版)"
            
        subtitle = doc.add_heading(f'ICT 練習題集{title_suffix}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # MCQ Section
        doc.add_heading('甲部 多項選擇題', level=1)
        doc.add_paragraph('請選擇最合適的答案。').bold = True
        
        # Process MCQ questions
        try:
            with open('MCQ.csv', 'r', encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file)
                question_counter = 1
                
                for row in reader:
                    if not row.get('QuestionText', '').strip():
                        continue
                    
                    # Extract source information
                    question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                    book_chapter = row.get('book_chapter', '')
                    topics = row.get('topics', '')
                    year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                    correct_answer = row.get('answer', '').upper() if preferences['show_answers'] else ''
                    
                    # Add source info if enabled
                    if preferences['show_source_info']:
                        source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                        source_para = doc.add_paragraph(source_info)
                        source_para.runs[0].font.size = Pt(10)
                        source_para.runs[0].italic = True
                    
                    # Add question
                    question_text = row['QuestionText']
                    marks = row['Marks']
                    question_para = doc.add_paragraph(f"{question_counter}. {question_text} ({marks} 分)")
                    question_para.runs[0].bold = True
                    
                    # Add options with answer highlighting
                    options = [('A', row['OptionA']), ('B', row['OptionB']), ('C', row['OptionC']), ('D', row['OptionD'])]
                    for letter, option_text in options:
                        option_para = doc.add_paragraph(f"   {letter}. {option_text}")
                        
                        # Highlight correct answer in red if showing answers
                        if preferences['show_answers'] and letter == correct_answer:
                            for run in option_para.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                run.bold = True
                            # Add checkmark
                            check_run = option_para.add_run(" ✓")
                            check_run.font.color.rgb = RGBColor(255, 0, 0)
                            check_run.bold = True
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Page break every 5 questions
                    if question_counter % 5 == 0:
                        doc.add_page_break()
                    
                    question_counter += 1
                    
        except FileNotFoundError:
            print("Error: 'MCQ.csv' not found.")
            return False
        except KeyError as e:
            print(f"Error: Missing column in MCQ.csv: {e}")
            return False
        
        # LQ Section
        doc.add_page_break()
        doc.add_heading('乙部 問答題', level=1)
        if preferences['show_answer_spaces']:
            doc.add_paragraph('請在適當的答案框內作答。').bold = True
        else:
            doc.add_paragraph('請回答以下問題。').bold = True
        
        # Process LQ questions
        try:
            with open('LQ.csv', 'r', encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file)
                lq_counter = 1
                
                for row in reader:
                    if not row.get('QuestionText', '').strip():
                        continue
                    
                    # Extract source information
                    question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                    book_chapter = row.get('book_chapter', '')
                    topics = row.get('topics', '')
                    year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                    
                    # Add source info if enabled
                    if preferences['show_source_info']:
                        source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                        source_para = doc.add_paragraph(source_info)
                        source_para.runs[0].font.size = Pt(10)
                        source_para.runs[0].italic = True
                    
                    # Add main question
                    question_text = row['QuestionText']
                    total_marks = row['Marks']
                    question_para = doc.add_paragraph(f"{lq_counter}. {question_text} ({total_marks} 分)")
                    question_para.runs[0].bold = True
                    
                    # Handle sub-questions
                    sub_questions = []
                    for i in range(1, 9):
                        sub_text = row.get(f'SubQuestion{i}', '') or ''
                        sub_marks = row.get(f'SubMarks{i}', '') or ''
                        sub_text = sub_text.strip() if sub_text else ''
                        sub_marks = sub_marks.strip() if sub_marks else ''
                        
                        if sub_text:
                            sub_questions.append({
                                'text': sub_text,
                                'marks': sub_marks,
                                'index': i
                            })
                    
                    # Display sub-questions
                    for idx, sub in enumerate(sub_questions):
                        sub_para = doc.add_paragraph(f"{sub['text']} ({sub['marks']}分)")
                        
                        # Show answers if enabled
                        if preferences['show_answers']:
                            # Look for answer in lq_answers
                            answer_text = ""
                            if question_id_with_years in lq_answers and sub['index'] in lq_answers[question_id_with_years]:
                                answer_text = lq_answers[question_id_with_years][sub['index']]
                            elif question_id_with_years in lq_answers and (idx + 1) in lq_answers[question_id_with_years]:
                                answer_text = lq_answers[question_id_with_years][idx + 1]
                            
                            if answer_text:
                                answer_para = doc.add_paragraph(f"答案: {answer_text}")
                                answer_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                answer_para.runs[0].bold = True
                            else:
                                answer_para = doc.add_paragraph("答案: [答案未提供]")
                                answer_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange color
                                answer_para.runs[0].italic = True
                        
                        # Show answer spaces if enabled
                        elif preferences['show_answer_spaces']:
                            marks_value = int(sub["marks"]) if sub["marks"] and sub["marks"].isdigit() else 2
                            lines = max(3, marks_value)  # At least 3 lines, more for higher marks
                            
                            # Add answer lines with proper spacing
                            for line_num in range(lines):
                                answer_line = doc.add_paragraph()
                                answer_line.add_run("_" * 60)  # Answer lines
                                answer_line.paragraph_format.space_before = Pt(6)
                                answer_line.paragraph_format.space_after = Pt(6)
                            
                            doc.add_paragraph()  # Extra spacing after answer area
                    
                    doc.add_page_break()  # New page for each LQ
                    lq_counter += 1
                    
        except FileNotFoundError:
            print("Note: 'LQ.csv' not found. Skipping long questions section.")
        except KeyError as e:
            print(f"Error: Missing column in LQ.csv: {e}")
        
        # Save the document
        doc.save('Exercise.docx')
        print("✓ DOCX file generated successfully!")
        return True
        
    except ImportError as e:
        print(f"✗ Cannot generate DOCX: {e}")
        print("Please install python-docx: pip3 install python-docx")
        return False
    except Exception as e:
        print(f"✗ Error creating Exercise.docx: {e}")
        return False

def main():
    """Main function to generate Markdown and DOCX formats with user preferences."""
    print("=== ICT Exercise Generator ===")
    
    # Get user preferences
    preferences = get_user_preferences()
    
    # Check and install packages for DOCX
    print("Checking required packages...")
    packages_installed = install_required_packages()
    
    print("\nGenerating exercise files...")
    
    # Generate Markdown (this always works)
    if generate_exercise_markdown(preferences):
        print("✓ Markdown file generated successfully")
    else:
        print("✗ Failed to generate Markdown file")
        return
    
    # Generate DOCX (only if packages are available)
    if packages_installed:
        try:
            if generate_exercise_docx(preferences):
                print("✓ DOCX file generated successfully")
            else:
                print("✗ Failed to generate DOCX file")
        except Exception as e:
            print(f"✗ DOCX generation failed: {e}")
    else:
        print("⚠ Skipping DOCX generation - python-docx not installed")
        print("To enable DOCX generation: pip3 install python-docx")
    
    print("\n=== Generation Complete ===")
    print("Files created:")
    
    # Check which files were actually created
    files_created = []
    if os.path.exists("Exercise.md"):
        files_created.append("- Exercise.md (Markdown format)")
    if os.path.exists("Exercise.docx"):
        files_created.append("- Exercise.docx (Word format)")
    
    if files_created:
        for file_info in files_created:
            print(file_info)
    else:
        print("No files were created successfully.")
    
    # Show configuration summary
    print(f"\nConfiguration used:")
    print(f"  • Source info: {'Shown' if preferences['show_source_info'] else 'Hidden'}")
    print(f"  • Answers: {'Shown' if preferences['show_answers'] else 'Hidden'}")
    print(f"  • Answer spaces: {'Shown' if preferences['show_answer_spaces'] else 'Hidden'}")
    
    # Show note about LQ answers if showing answers but no answer file found
    if preferences['show_answers']:
        if not os.path.exists("LQ_Answers.csv"):
            print("\nNote: Create 'LQ_Answers.csv' file with long question answers for better results.")
            print("Format: QuestionIDwithQuestionyears,SubAnswer1,SubAnswer2,SubAnswer3,...")
    
    # Show tip about converting DOCX to PDF if needed
    if os.path.exists("Exercise.docx"):
        print("\nTip: To convert DOCX to PDF:")
        print("  • Open Exercise.docx in Word/Pages and save as PDF")
        print("  • Use online converters like SmallPDF or ILovePDF")

# Run the script
if __name__ == "__main__":
    main()