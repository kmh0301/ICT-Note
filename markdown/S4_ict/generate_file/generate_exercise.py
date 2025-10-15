import csv
import subprocess
import sys
import os
import glob

CSV_FOLDER = '/Users/sallypang/Documents/GitHub/ICT-Note/markdown/S4_ict/generate_file/csv'
RESULT_FOLDER = '/Users/sallypang/Documents/GitHub/ICT-Note/markdown/S4_ict/generate_file/result'

def choose_csv_file(prompt="Select CSV file"):
    csv_files = glob.glob(os.path.join(CSV_FOLDER, '*.csv'))
    csv_names = [os.path.basename(f) for f in csv_files]
    if not csv_names:
        print("No CSV files found in the folder.")
        sys.exit(1)
    print(f"\n{prompt}:")
    for idx, fname in enumerate(csv_names, 1):
        print(f"  {idx}. {fname}")
    while True:
        choice = input("Enter file name (e.g., MCQ.csv): ").strip()
        if choice in csv_names:
            return os.path.join(CSV_FOLDER, choice)
        print("Not found - please type one of the names above.")

def load_lq_answers(lq_answers_path):
    lq_answers = {}
    if not lq_answers_path or not os.path.exists(lq_answers_path):
        print("Note: LQ_Answers.csv not found. Long question answers will not be displayed.")
        return lq_answers
    try:
        with open(lq_answers_path, 'r', encoding='utf-8') as csv_file:
            reader = csv.DictReader(csv_file)
            for row in reader:
                question_id = row.get('QuestionIDwithQuestionyears', '')
                if question_id:
                    answers = {}
                    for i in range(1, 9):
                        answer_key = f'SubAnswer{i}'
                        if answer_key in row and row[answer_key].strip():
                            answers[i] = row[answer_key].strip()
                    if answers:
                        lq_answers[question_id] = answers
    except Exception as e:
        print(f"Error reading LQ answers: {e}")
    return lq_answers

def get_user_preferences():
    print("\n=== ICT Exercise Generator Configuration ===")
    print("Please choose what to include in your exercise files:\n")
    while True:
        show_source = input("Show source information (year, chapter, topic)? (y/n): ").lower().strip()
        if show_source in ['y', 'yes', 'n', 'no']:
            show_source_info = show_source in ['y', 'yes']
            break
        print("Please enter 'y' for yes or 'n' for no")
    while True:
        show_answer = input("Show answers? (y/n): ").lower().strip()
        if show_answer in ['y', 'yes', 'n', 'no']:
            show_answers = show_answer in ['y', 'yes']
            break
        print("Please enter 'y' for yes or 'n' for no")
    if not show_answers:
        while True:
            show_space = input("Show answer spaces for long questions? (y/n): ").lower().strip()
            if show_space in ['y', 'yes', 'n', 'no']:
                show_answer_spaces = show_space in ['y', 'yes']
                break
            print("Please enter 'y' for yes or 'n' for no")
    else:
        show_answer_spaces = False
    preferences = {
        'show_source_info': show_source_info,
        'show_answers': show_answers,
        'show_answer_spaces': show_answer_spaces
    }
    print(f"\n📋 Your preferences:")
    print(f"  • Source info: {'Yes' if show_source_info else 'No'}")
    print(f"  • Show answers: {'Yes' if show_answers else 'No'}")
    print(f"  • Answer spaces: {'Yes' if show_answer_spaces else 'No'}")
    print()
    return preferences

def install_required_packages():
    packages_to_install = []
    try:
        import docx
        print("✓ python-docx is already installed")
    except ImportError:
        packages_to_install.append('python-docx')
    for package in packages_to_install:
        try:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"✓ Successfully installed {package}")
        except subprocess.CalledProcessError as e:
            print(f"✗ Failed to install {package}")
            print("Please try installing manually:")
            print(f"  pip3 install {package}")
            print(f"  or: python3 -m pip install {package}")
            return False
    return True

def generate_exercise_markdown(preferences, mcq_csv_path, lq_csv_path, lq_answers_path):
    import os
    lq_answers = {}
    if preferences['show_answers']:
        lq_answers = load_lq_answers(lq_answers_path)
    markdown_path = os.path.join(RESULT_FOLDER, 'Exercise.md')
    try:
        with open(markdown_path, 'w', encoding='utf-8') as md_file:
            # --- YAML frontmatter ---
            md_file.write("---\n")
            md_file.write("marp: true\n")
            md_file.write("theme: UT\n")
            md_file.write("class: title-page\n")
            md_file.write("paginate: false\n")
            md_file.write("backgroundColor: white\n")
            md_file.write("---\n\n")

            # --- Title page ---
            title_suffix = ""
            if preferences['show_answers']:
                title_suffix = " (答案版)"
            elif not preferences['show_answer_spaces']:
                title_suffix = " (純題目版)"
            md_file.write("# **佛教黃鳳翎中學**\n\n")
            md_file.write(f"# **ICT 練習題集{title_suffix}**\n\n")
            md_file.write("---\n\n")

            # --- MCQ Section ---
            md_file.write('<div class="section-title"><strong>甲部 多項選擇題</strong></div>\n')
            md_file.write('<strong>請選擇最合適的答案。</strong>\n\n')
            try:
                with open(mcq_csv_path, 'r', encoding='utf-8') as csv_file:
                    reader = csv.DictReader(csv_file)
                    question_counter = 1
                    for row in reader:
                        if not row.get('QuestionText', '').strip():
                            continue
                        question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                        book_chapter = row.get('book_chapter', '')
                        topics = row.get('topics', '')
                        year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                        question_text = row['QuestionText']
                        option_a = row['OptionA']
                        option_b = row['OptionB']
                        option_c = row['OptionC']
                        option_d = row['OptionD']
                        marks = row['Marks']
                        correct_answer = row.get('answer', '').upper() if preferences['show_answers'] else ''
                        # Add a divider after every 5 questions
                        if question_counter > 1 and question_counter % 5 == 1:
                            md_file.write('---\n\n')
                        # Source info
                        if preferences['show_source_info']:
                            source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                            md_file.write(f'<div class="question-source">{source_info}</div><br>\n')
                        md_file.write(f'<div class="question-item">{question_counter}. {question_text} <span class="points">({marks} 分)</span></div>\n')
                        md_file.write('<ul class="mcq-options">\n')
                        # Write options
                        options = [('A', option_a), ('B', option_b), ('C', option_c), ('D', option_d)]
                        for letter, option_text in options:
                            if preferences['show_answers'] and letter == correct_answer:
                                md_file.write(f'<li style="color: red; font-weight: bold;">{option_text} ✓</li>\n')
                            else:
                                md_file.write(f'<li>{option_text}</li>\n')
                        md_file.write('</ul>\n\n')
                        question_counter += 1
            except FileNotFoundError:
                print(f"Error: File {mcq_csv_path} not found.")
                return False
            except KeyError as e:
                print(f"Error: Missing column in MCQ file: {e}")
                return False

            # --- LQ Section ---
            md_file.write('---\n\n')
            md_file.write('<div class="section-title"><strong>乙部 問答題</strong></div>\n')
            if preferences['show_answer_spaces']:
                md_file.write('<strong>請在適當的答案框內作答。</strong>\n\n')
            else:
                md_file.write('<strong>請回答以下問題。</strong>\n\n')
            try:
                with open(lq_csv_path, 'r', encoding='utf-8') as csv_file:
                    reader = csv.DictReader(csv_file)
                    lq_counter = 1
                    for row in reader:
                        if not row.get('QuestionText', '').strip():
                            continue
                        question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                        book_chapter = row.get('book_chapter', '')
                        topics = row.get('topics', '')
                        year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                        question_id = question_id_with_years.split('-')[-1] if question_id_with_years else str(lq_counter)
                        question_text = row['QuestionText']
                        total_marks = row['Marks']
                        md_file.write('---\n\n')
                        # Source info
                        if preferences['show_source_info']:
                            source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                            md_file.write(f'<div class="question-source">{source_info}</div><br>\n')
                        md_file.write(f'<div class="question-item">{lq_counter}. {question_text} <span class="points">({total_marks} 分)</span></div>\n\n')
                        # Handle sub-questions
                        sub_questions = []
                        for i in range(1, 9):  # SubQuestion1 to SubQuestion8
                            sub_text = row.get(f'SubQuestion{i}', '') or ''
                            sub_marks = row.get(f'SubMarks{i}', '') or ''
                            sub_text = sub_text.strip() if sub_text else ''
                            sub_marks = sub_marks.strip() if sub_marks else ''
                            if sub_text:
                                sub_questions.append({
                                    'text': sub_text,
                                    'marks': sub_marks,
                                    'index': i
                                })
                        # Display sub-questions
                        for idx, sub in enumerate(sub_questions):
                            md_file.write(f'<div class="sub-question-item">{sub["text"]} <span class="points">({sub["marks"]}分)</span></div>\n')
                            # Show answers if enabled
                            if preferences['show_answers']:
                                answer_text = ""
                                if question_id_with_years in lq_answers and sub['index'] in lq_answers[question_id_with_years]:
                                    answer_text = lq_answers[question_id_with_years][sub['index']]
                                elif question_id_with_years in lq_answers and (idx + 1) in lq_answers[question_id_with_years]:
                                    answer_text = lq_answers[question_id_with_years][idx + 1]
                                if answer_text:
                                    md_file.write(f'<div style="color: red; font-weight: bold; margin: 10px 0; padding: 10px; border-left: 3px solid red;">答案: {answer_text}</div>\n')
                                else:
                                    md_file.write('<div style="color: orange; font-style: italic;">答案: [答案未提供]</div>\n')
                            elif preferences['show_answer_spaces']:
                                marks_value = int(sub["marks"]) if sub["marks"] and sub["marks"].isdigit() else 2
                                if marks_value >= 4:
                                    height = "5em"
                                elif marks_value >= 2:
                                    height = "4em"
                                else:
                                    height = "3em"
                                md_file.write(f'<div class="answer-box"><div style="height: {height};"></div></div>\n')
                            md_file.write('\n')
                        lq_counter += 1
            except FileNotFoundError:
                print(f"Note: File {lq_csv_path} not found. Skipping long questions section.")
            except KeyError as e:
                print(f"Error: Missing column in LQ file: {e}")
        print("Successfully generated Exercise.md!")
        return True
    except Exception as e:
        print(f"Error creating Exercise.md: {e}")
        return False
def generate_exercise_docx(preferences, mcq_csv_path, lq_csv_path, lq_answers_path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os

    lq_answers = {}
    if preferences['show_answers']:
        lq_answers = load_lq_answers(lq_answers_path)
    docx_path = os.path.join(RESULT_FOLDER, 'Exercise.docx')
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # Title page
        title_suffix = ""
        if preferences['show_answers']:
            title_suffix = " (答案版)"
        elif not preferences['show_answer_spaces']:
            title_suffix = " (純題目版)"
        title = doc.add_heading('佛教黃鳳翎中學', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading(f'ICT 練習題集{title_suffix}', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # MCQ Section
        doc.add_heading('甲部 多項選擇題', level=1)
        doc.add_paragraph('請選擇最合適的答案。').bold = True
        try:
            with open(mcq_csv_path, 'r', encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file)
                question_counter = 1
                for row in reader:
                    if not row.get('QuestionText', '').strip():
                        continue
                    question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                    book_chapter = row.get('book_chapter', '')
                    topics = row.get('topics', '')
                    year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                    correct_answer = row.get('answer', '').upper() if preferences['show_answers'] else ''
                    if preferences['show_source_info']:
                        source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                        source_para = doc.add_paragraph(source_info)
                        source_para.runs[0].font.size = Pt(10)
                        source_para.runs[0].italic = True
                    question_text = row['QuestionText']
                    marks = row['Marks']
                    question_para = doc.add_paragraph(f"{question_counter}. {question_text} ({marks} 分)")
                    question_para.runs[0].bold = True
                    options = [('A', row['OptionA']), ('B', row['OptionB']), ('C', row['OptionC']), ('D', row['OptionD'])]
                    for letter, option_text in options:
                        option_para = doc.add_paragraph(f"   {letter}. {option_text}")
                        if preferences['show_answers'] and letter == correct_answer:
                            for run in option_para.runs:
                                run.font.color.rgb = RGBColor(255, 0, 0)
                                run.bold = True
                            check_run = option_para.add_run(" ✓")
                            check_run.font.color.rgb = RGBColor(255, 0, 0)
                            check_run.bold = True
                    doc.add_paragraph()
                    if question_counter % 5 == 0:
                        doc.add_page_break()
                    question_counter += 1
        except FileNotFoundError:
            print(f"Error: MCQ file not found.")
            return False
        except KeyError as e:
            print(f"Error: Missing column in MCQ file: {e}")
            return False

        doc.add_page_break()
        doc.add_heading('乙部 問答題', level=1)
        if preferences['show_answer_spaces']:
            doc.add_paragraph('請在適當的答案框內作答。').bold = True
        else:
            doc.add_paragraph('請回答以下問題。').bold = True
        try:
            with open(lq_csv_path, 'r', encoding='utf-8') as csv_file:
                reader = csv.DictReader(csv_file)
                lq_counter = 1
                for row in reader:
                    if not row.get('QuestionText', '').strip():
                        continue
                    question_id_with_years = row.get('QuestionIDwithQuestionyears', '')
                    book_chapter = row.get('book_chapter', '')
                    topics = row.get('topics', '')
                    year = question_id_with_years.split('-')[0] if question_id_with_years else ''
                    if preferences['show_source_info']:
                        source_info = f"({year}_{question_id_with_years}_{book_chapter}_{topics})"
                        source_para = doc.add_paragraph(source_info)
                        source_para.runs[0].font.size = Pt(10)
                        source_para.runs[0].italic = True
                    question_text = row['QuestionText']
                    total_marks = row['Marks']
                    question_para = doc.add_paragraph(f"{lq_counter}. {question_text} ({total_marks} 分)")
                    question_para.runs[0].bold = True
                    sub_questions = []
                    for i in range(1,9):
                        sub_text = row.get(f'SubQuestion{i}', '') or ''
                        sub_marks = row.get(f'SubMarks{i}', '') or ''
                        sub_text = sub_text.strip() if sub_text else ''
                        sub_marks = sub_marks.strip() if sub_marks else ''
                        if sub_text:
                            sub_questions.append({
                                'text': sub_text,
                                'marks': sub_marks,
                                'index': i
                            })
                    for idx, sub in enumerate(sub_questions):
                        sub_para = doc.add_paragraph(f"{sub['text']} ({sub['marks']}分)")
                        if preferences['show_answers']:
                            answer_text = ""
                            if question_id_with_years in lq_answers and sub['index'] in lq_answers[question_id_with_years]:
                                answer_text = lq_answers[question_id_with_years][sub['index']]
                            elif question_id_with_years in lq_answers and (idx+1) in lq_answers[question_id_with_years]:
                                answer_text = lq_answers[question_id_with_years][idx+1]
                            if answer_text:
                                answer_para = doc.add_paragraph(f"答案: {answer_text}")
                                answer_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                                answer_para.runs[0].bold = True
                            else:
                                answer_para = doc.add_paragraph("答案: [答案未提供]")
                                answer_para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
                                answer_para.runs[0].italic = True
                        elif preferences['show_answer_spaces']:
                            marks_value = int(sub["marks"]) if sub["marks"] and sub["marks"].isdigit() else 2
                            lines = max(3, marks_value)
                            for line_num in range(lines):
                                answer_line = doc.add_paragraph()
                                answer_line.add_run("_" * 60)
                                answer_line.paragraph_format.space_before = Pt(6)
                                answer_line.paragraph_format.space_after = Pt(6)
                            doc.add_paragraph()
                    doc.add_page_break()
                    lq_counter += 1
        except FileNotFoundError:
            print(f"Note: LQ file not found. Skipping long questions section.")
        doc.save(docx_path)
        print("✓ DOCX file generated successfully!")
        return True
    except Exception as e:
        print(f"✗ Error creating Exercise.docx: {e}")
        return False

def main():
    print("=== ICT Exercise Generator ===")
    preferences = get_user_preferences()
    print("\nChoosing MCQ and LQ CSV files:")
    mcq_csv_path = choose_csv_file("Select MCQ csv file")
    lq_csv_path = choose_csv_file("Select LQ csv file")
    lq_answers_path = os.path.join(CSV_FOLDER, 'LQ_Answers.csv')

    print("\nGenerating exercise files...")
    result_md = generate_exercise_markdown(preferences, mcq_csv_path, lq_csv_path, lq_answers_path)
    print("✓ Markdown file generated successfully" if result_md else "✗ Failed to generate Markdown file")

    # DOCX
    try:
        import docx
        result_docx = generate_exercise_docx(preferences, mcq_csv_path, lq_csv_path, lq_answers_path)
        print("✓ DOCX file generated successfully" if result_docx else "✗ Failed to generate DOCX file")
    except ImportError:
        print("⚠ DOCX file not created: python-docx not installed. Run pip3 install python-docx first.")

    print("\n=== Generation Complete ===")
    markdown_path = os.path.join(RESULT_FOLDER, "Exercise.md")
    docx_path = os.path.join(RESULT_FOLDER, "Exercise.docx")
    files_created = []
    if os.path.exists(markdown_path):
        files_created.append("- Exercise.md (Markdown format)")
    if os.path.exists(docx_path):
        files_created.append("- Exercise.docx (Word format)")
    print('\n'.join(files_created) if files_created else "No files were created successfully.")

if __name__ == "__main__":
    main()
