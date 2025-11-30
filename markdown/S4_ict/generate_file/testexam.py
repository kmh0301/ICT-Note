import json
import subprocess
import sys
import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURATION ---
# Rename CSV_FOLDER to DATA_FOLDER since we use JSON now
DATA_FOLDER = '/Users/sallypang/Documents/GitHub/ICT-Note/markdown/S4_ict/generate_file/json' 
RESULT_FOLDER = '/Users/sallypang/Documents/GitHub/ICT-Note/markdown/S4_ict/generate_file/result'
# --- END CONFIGURATION ---

# --- Utility Functions ---
def choose_json_file(prompt="Select JSON file"):
    """Lists JSON files in the folder and asks user to pick one."""
    if not os.path.exists(DATA_FOLDER):
        print(f"Error: Folder not found: {DATA_FOLDER}")
        sys.exit(1)
        
    json_files = glob.glob(os.path.join(DATA_FOLDER, '*.json'))
    json_names = [os.path.basename(f) for f in json_files]
    
    if not json_names:
        print(f"No JSON files found in {DATA_FOLDER}.")
        sys.exit(1)
        
    print(f"\n{prompt}:")
    for idx, fname in enumerate(json_names, 1):
        print(f"  {idx}. {fname}")
        
    while True:
        choice = input("Enter file name (e.g., exam_data.json) or number: ").strip()
        # Handle number input
        if choice.isdigit():
            idx = int(choice) - 1
            if 0 <= idx < len(json_names):
                return os.path.join(DATA_FOLDER, json_names[idx])
        # Handle name input
        if choice in json_names:
            return os.path.join(DATA_FOLDER, choice)
        print("Not found - please try again.")

def load_exam_data(json_path):
    """Loads the JSON file and splits it into MC and LQ lists."""
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        mc_questions = [q for q in data if q.get('type') == 'MC']
        lq_questions = [q for q in data if q.get('type') == 'LQ']
        
        return mc_questions, lq_questions
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)

def install_required_packages():
    packages_to_install = []
    try:
        import docx
        print("✓ python-docx is already installed")
    except ImportError:
        packages_to_install.append('python-docx')
    for package in packages_to_install:
        try:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"✓ Successfully installed {package}")
        except subprocess.CalledProcessError:
            print(f"✗ Failed to install {package}")
            print("Please try installing manually: pip3 install python-docx")
            return False
    return True

def get_user_preferences():
    """Gets user configuration for exercise generation."""
    print("\n=== ICT Exercise Generator (JSON Version) ===")
    
    # 1. Output Preferences
    print("\nPlease choose what to include:")
    show_source = input("Show source info (year/topic)? (y/n) [n]: ").lower().strip() == 'y'
    show_answers = input("Show answers? (y/n) [n]: ").lower().strip() == 'y'
    
    show_answer_spaces = False
    if not show_answers:
        show_answer_spaces = input("Show answer spaces for long questions? (y/n) [y]: ").lower().strip() != 'n'
        
    # 2. Exam Details
    print("\n--- Exam Paper Details ---")
    exam_duration = input("Enter Exam Duration (e.g., 40 分鐘) [40 分鐘]: ").strip() or "40 分鐘"
    exam_date = input("Enter Exam Date (e.g., 2026年2月29日): ").strip() or "2026年2月29日"
    total_marks = input("Enter Total Marks (e.g., 50 分) [50 分]: ").strip() or "50 分"
        
    return {
        'show_source_info': show_source,
        'show_answers': show_answers,
        'show_answer_spaces': show_answer_spaces,
        'exam_duration': f"考試時間：{exam_duration}",
        'exam_date': f"考試日期：{exam_date}",
        'total_marks': f"總分: {total_marks}"
    }
# --- Sub-Question Auto-numbering and Cleaning Logic ---
def get_auto_numbered_sub_text(sub_text, sub_idx):
    """
    Cleans the sub-question text from existing complex numbering (e.g., (a), (b)(i)) 
    and prepends the new index-based alphabetical numbering (a), (b), etc.
    """
    
    # 1. Generate the new prefix: (a), (b), (c)...
    prefix_char = chr(97 + sub_idx) # 97 is ASCII for 'a'
    new_prefix = f'({prefix_char}) '
    
    cleaned_text = sub_text.strip()
    
    # 2. Attempt to strip complex existing prefixes (e.g., "(a)", "(b) (i)")
    # We look for something starting with '(' and containing ')' near the start.
    if cleaned_text.startswith('('):
        # Find the end of the *first* parenthetical group
        end_paren_index = cleaned_text.find(')')
        
        if end_paren_index > 0 and end_paren_index < 10: 
            # Check if the text continues after the parenthesis (e.g., space or text)
            potential_rest_of_text = cleaned_text[end_paren_index + 1:].lstrip()
            
            # If the next characters are also part of a prefix like ` (i) ` or just space, 
            # skip those too to get to the actual question text.
            if potential_rest_of_text.startswith('(') and ')' in potential_rest_of_text:
                # Handle complex cases like "(b) (i)"
                second_end_paren = potential_rest_of_text.find(')')
                if second_end_paren > 0:
                    cleaned_text = potential_rest_of_text[second_end_paren + 1:].lstrip()
                else:
                    cleaned_text = potential_rest_of_text
            else:
                # Handle simple cases like "(a)"
                cleaned_text = potential_rest_of_text
        
    # 3. Prepend the new, cleaned prefix
    return f'{new_prefix}{cleaned_text}'


# --- CORE FUNCTION 1: Generate Markdown ---
def generate_exercise_markdown(preferences, mc_questions, lq_questions):
    if not os.path.exists(RESULT_FOLDER):
        os.makedirs(RESULT_FOLDER)
        
    markdown_path = os.path.join(RESULT_FOLDER, 'Exercise.md')
    
    # Configuration
    SCHOOL_NAME = "佛教黃鳳翎中學"
    EXAM_TITLE = "2025/2026 上學期考試"
    SUBJECT = "中四級資訊及通訊科技"
    
    try:
        with open(markdown_path, 'w', encoding='utf-8') as md_file:
            # --- YAML Frontmatter ---
            md_file.write("---\nmarp: true\ntheme: testexam\nclass: title-page\npaginate: false\nbackgroundColor: white\n---\n\n")

            # --- Slide 1: Front Page ---
            title_suffix = " (答案版)" if preferences['show_answers'] else ""
            
            md_file.write(f'<div class="school-name">{SCHOOL_NAME}</div>\n') 
            md_file.write(f'<div class="exam-title-main">{EXAM_TITLE}{title_suffix}</div>\n') 
            md_file.write(f'<div class="subject-title">{SUBJECT}</div>\n') 
            md_file.write(f'<div class="answer-book">試題答題簿</div>\n\n') 
            
            # Student Info & Details
            md_file.write('<div class="student-info-block">\n')
            md_file.write('  <p class="info-line">班別: _______________</p>\n') 
            md_file.write('  <p class="info-line">班號: _______________</p>\n')
            md_file.write('  <p class="info-line">姓名: _______________</p>\n')
            md_file.write('</div>\n\n')

            md_file.write('<div class="date-marks-block">\n')
            md_file.write(f'  <p class="date-line">{preferences["exam_date"]}</p>\n')
            md_file.write(f'  <p class="date-line">{preferences["exam_duration"]}</p>\n')
            md_file.write('  <table><tr>\n')
            md_file.write(f'    <td class="total-marks-label">{preferences["total_marks"]}</td>\n')
            md_file.write('    <td class="score-box"></td>\n')
            md_file.write('  </tr></table>\n')
            md_file.write('</div>\n\n')
            
            md_file.write("---\n\n") 
            
            # --- SECTION A: MCQ ---
            if mc_questions:
                md_file.write('<div class="section-title">**甲部 多項選擇題**</div>\n')
                md_file.write('**請選擇最合適的答案。**\n\n')
                
                for idx, q in enumerate(mc_questions, 1):
                    if idx > 1: md_file.write('---\n')
                    
                    if preferences['show_source_info']:
                        source_info = f"({q.get('id', '')}_{q.get('book_chapter', '')}_{q.get('topics', '')})"
                        md_file.write(f'<p style="font-size: 0.7em; color: gray;">{source_info}</p>\n')
                    
                    md_file.write(f'## **{idx}. {q["text"]}** <span class="points">({q["marks"]} 分)</span>\n')
                    
                    md_file.write('<ul>\n')
                    options = q.get('options', {})
                    correct_ans = q.get('answer', '').upper()
                    
                    for letter in ['A', 'B', 'C', 'D']:
                        opt_text = options.get(letter, '')
                        if preferences['show_answers'] and letter == correct_ans:
                            md_file.write(f'<li style="color: red; font-weight: bold;">{letter}. {opt_text} ✓</li>\n')
                        else:
                            md_file.write(f'<li>{letter}. {opt_text}</li>\n')
                    md_file.write('</ul>\n\n')

            # --- SECTION B: LQ ---
            if lq_questions:
                md_file.write('---\n\n<div class="section-title">**乙部 問答題**</div>\n')
                md_file.write('請在適當的答案框內作答。\n\n' if preferences['show_answer_spaces'] else '請回答以下問題。\n\n')
                
                for idx, q in enumerate(lq_questions, 1):
                    
                    if preferences['show_source_info']:
                        source_info = f"({q.get('id', '')}_{q.get('book_chapter', '')})"
                        md_file.write(f'<p style="font-size: 0.7em; color: gray;">{source_info}</p>\n')
                        
                    md_file.write(f'## **{idx}. {q["text"]}** <span class="points">({q["marks"]} 分)</span>\n\n')
                    
                    # Dynamic Sub-questions handling (Auto-numbering implemented here)
                    sub_qs = q.get('sub_questions', [])
                    for sub_idx, sub in enumerate(sub_qs):
                        
                        # Apply auto-numbering and cleaning logic
                        full_sub_text = get_auto_numbered_sub_text(sub.get('text', ''), sub_idx)
                        
                        sub_marks = sub.get('marks', '')
                        sub_ans = sub.get('answer', '')
                        
                        md_file.write(f'### **{full_sub_text}** <span class="points">({sub_marks}分)</span>\n')
                        
                        if preferences['show_answers']:
                            if sub_ans:
                                md_file.write(f'<div style="color: red; font-weight: bold; border-left: 3px solid red; padding-left: 10px;">答案: {sub_ans}</div>\n')
                            else:
                                md_file.write('<div style="color: orange;">[未提供答案]</div>\n')
                        elif preferences['show_answer_spaces']:
                            # Estimate space based on marks
                            try:
                                m_val = int(sub_marks)
                            except:
                                m_val = 2
                            height = "5em" if m_val >= 4 else "3em"
                            md_file.write(f'<div style="border: 1px solid #ccc; margin: 5px 0 15px 0;"><div style="height: {height};"></div></div>\n')
                        
                        md_file.write('---\n')

        print("✓ Exercise.md generated successfully")
        return True
    except Exception as e:
        print(f"Error creating Markdown: {e}")
        return False

# --- CORE FUNCTION 2: Generate DOCX ---
def generate_exercise_docx(preferences, mc_questions, lq_questions):
    if not os.path.exists(RESULT_FOLDER):
        os.makedirs(RESULT_FOLDER)
        
    docx_path = os.path.join(RESULT_FOLDER, 'Exercise.docx')
    SCHOOL_NAME = "佛教黃鳳翎中學"
    
    try:
        doc = Document()
        # Setup styles/margins... (kept basic for brevity)
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Header
        title = doc.add_heading(SCHOOL_NAME, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_text = "ICT 練習題集" + (" (答案版)" if preferences['show_answers'] else "")
        subtitle = doc.add_heading(subtitle_text, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = f"{preferences['exam_date']} | {preferences['exam_duration']} | {preferences['total_marks']}"
        doc.add_paragraph(info).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # --- MCQ ---
        if mc_questions:
            doc.add_heading('甲部 多項選擇題', level=1)
            doc.add_paragraph('請選擇最合適的答案。').bold = True
            
            for idx, q in enumerate(mc_questions, 1):
                # Source info
                if preferences['show_source_info']:
                    source_txt = f"({q.get('id', '')} {q.get('topics', '')})"
                    p = doc.add_paragraph(source_txt)
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # Question
                p = doc.add_paragraph(f"{idx}. {q['text']} ({q['marks']} 分)")
                p.runs[0].bold = True
                
                # Options
                options = q.get('options', {})
                correct_ans = q.get('answer', '').upper()
                
                for letter in ['A', 'B', 'C', 'D']:
                    opt_text = options.get(letter, '')
                    p_opt = doc.add_paragraph(f"   {letter}. {opt_text}")
                    
                    if preferences['show_answers'] and letter == correct_ans:
                        for run in p_opt.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.bold = True
                        p_opt.add_run(" ✓").font.color.rgb = RGBColor(255, 0, 0)
                
                doc.add_paragraph() # Spacer
                # Removed page break here as it can make DOCX too fragmented, let it flow naturally

        doc.add_page_break()

        # --- LQ ---
        if lq_questions:
            doc.add_heading('乙部 問答題', level=1)
            doc.add_paragraph('請在適當的答案框內作答。' if preferences['show_answer_spaces'] else '請回答以下問題。').bold = True
            
            for idx, q in enumerate(lq_questions, 1):
                if preferences['show_source_info']:
                    source_txt = f"({q.get('id', '')})"
                    p = doc.add_paragraph(source_txt)
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                    
                p = doc.add_paragraph(f"{idx}. {q['text']} ({q['marks']} 分)")
                p.runs[0].bold = True
                
                # Sub-questions Loop (Auto-numbering implemented here)
                for sub_idx, sub in enumerate(q.get('sub_questions', [])):
                    
                    # Apply auto-numbering and cleaning logic
                    full_sub_text = get_auto_numbered_sub_text(sub.get('text', ''), sub_idx)
                    
                    # Add sub-question text
                    p_sub = doc.add_paragraph(f"{full_sub_text} ({sub.get('marks', '')}分)")
                    
                    if preferences['show_answers']:
                        ans = sub.get('answer', '[答案未提供]')
                        p_ans = doc.add_paragraph(f"答案: {ans}")
                        p_ans.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    elif preferences['show_answer_spaces']:
                        # Draw lines
                        try:
                            lines = max(3, int(sub['marks']))
                        except:
                            lines = 3
                        for _ in range(lines):
                            p_line = doc.add_paragraph("_" * 60)
                            p_line.paragraph_format.line_spacing = Pt(18)
                
                doc.add_page_break()

        doc.save(docx_path)
        print("✓ Exercise.docx generated successfully")
        return True
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return False
    
# --- Main Execution ---
def main():
    install_required_packages()
    
    # 1. Select JSON File (One file containing both MC and LQ)
    json_path = choose_json_file()
    
    # 2. Load and parse Data
    mc_data, lq_data = load_exam_data(json_path)
    print(f"Loaded {len(mc_data)} MCQs and {len(lq_data)} LQs.")
    
    # 3. Get Preferences
    prefs = get_user_preferences()
    
    print("\nGenerating files...")
    # 4. Generate
    generate_exercise_markdown(prefs, mc_data, lq_data)
    generate_exercise_docx(prefs, mc_data, lq_data)
    
    print("\n=== Generation Complete ===")
    print(f"Files saved in: {RESULT_FOLDER}")

if __name__ == "__main__":
    main()