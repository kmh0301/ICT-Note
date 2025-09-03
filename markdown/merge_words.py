from docx import Document
import os
import re
from pathlib import Path

def merge_word_documents_chapter_order(input_folder, output_filename):
    """
    Merge Word documents in chapter/checkpoint order (C-X_Checkpoint_CE_Y_TE.docx)
    """
    # Create a new document for the merged content
    merged_doc = Document()
    
    # Get all .docx files from the input folder
    input_path = Path(input_folder)
    word_files = list(input_path.glob("*.docx"))
    
    if not word_files:
        print(f"No Word documents found in {input_folder}")
        return
    
    # Sort by chapter and checkpoint numbers
    def extract_chapter_checkpoint(filename):
        # Extract chapter number (after C-) and checkpoint number (after CE_)
        match = re.search(r'C-(\d+).*CE_(\d+)', filename.name)
        if match:
            chapter = int(match.group(1))
            checkpoint = int(match.group(2))
            return (chapter, checkpoint)
        return (999, 999)  # Put unmatched files at the end
    
    word_files.sort(key=extract_chapter_checkpoint)
    
    print(f"\nFound {len(word_files)} Word documents to merge in chapter/checkpoint order:")
    for i, file in enumerate(word_files, 1):
        chapter_info = extract_chapter_checkpoint(file)
        print(f"  {i:2d}. {file.name} (Chapter {chapter_info[0]}, Checkpoint {chapter_info[1]})")
    
    # Process each Word document
    for i, file_path in enumerate(word_files):
        try:
            print(f"Processing: {file_path.name}")
            
            # Open the current document
            current_doc = Document(file_path)
            
            # Add a page break before each new document (except the first one)
            if i > 0:
                merged_doc.add_page_break()
            
            # Add document title
            title_paragraph = merged_doc.add_paragraph()
            title_run = title_paragraph.add_run(f"Document: {file_path.name}")
            title_run.bold = True
            merged_doc.add_paragraph()  # Add empty line
            
            # Copy all paragraphs from current document
            for paragraph in current_doc.paragraphs:
                new_paragraph = merged_doc.add_paragraph()
                
                # Copy paragraph style and content
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    # Copy formatting
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
                    if run.font.name:
                        new_run.font.name = run.font.name
            
            # Copy tables if any
            for table in current_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
            
        except Exception as e:
            print(f"Error processing {file_path.name}: {str(e)}")
            continue
    
    # Save the merged document
    try:
        merged_doc.save(output_filename)
        print(f"\nMerged document saved as: {output_filename}")
        print("Merge completed successfully!")
    except Exception as e:
        print(f"Error saving merged document: {str(e)}")

if __name__ == "__main__":
    # Simple usage - automatically sorts by chapter/checkpoint order
    input_folder = "./documents"
    output_filename = "merged_checkpoint_ordered.docx"
    
    print("=== Word Document Merger - Chapter Order ===")
    print(f"Looking for Word documents in: {input_folder}")
    print(f"Output will be saved as: {output_filename}")
    
    merge_word_documents_chapter_order(input_folder, output_filename)